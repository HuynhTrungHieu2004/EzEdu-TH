#!/usr/bin/env python3
"""Build the polished Vietnamese system-analysis DOCX from the Markdown source."""

from __future__ import annotations

from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Bao_cao_phan_tich_nghiep_vu_EzEdu_AI.md"
OUTPUT = ROOT / "Bao_cao_phan_tich_nghiep_vu_EzEdu_AI.docx"

NAVY = "163B65"
BLUE = "2E74B5"
BLUE_LIGHT = "DCEAF7"
INK = "182230"
MUTED = "526173"
LINE = "C8D4E1"
TABLE_HEADER = "F2F4F7"
PALE = "F4F7FB"
WHITE = "FFFFFF"

LETTER_W = Inches(8.5)
LETTER_H = Inches(11)
PORTRAIT_WIDTH = Inches(6.5)
LANDSCAPE_WIDTH = Inches(9.0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_indent(table, indent_twips=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text_value: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text_value
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_picture_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def apply_run_style(run, *, bold=False, italic=False, code=False, color=None, size=None) -> None:
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("7A3E00")
    else:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def add_inline(paragraph, value: str, *, base_bold=False, base_color=None, base_size=None) -> None:
    value = value.replace("\\|", "|")
    pos = 0
    for match in TOKEN_RE.finditer(value):
        if match.start() > pos:
            run = paragraph.add_run(value[pos : match.start()])
            apply_run_style(run, bold=base_bold, color=base_color, size=base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            apply_run_style(run, bold=True, color=base_color, size=base_size)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            apply_run_style(run, code=True)
        else:
            run = paragraph.add_run(token[1:-1])
            apply_run_style(run, italic=True, color=base_color, size=base_size)
        pos = match.end()
    if pos < len(value):
        run = paragraph.add_run(value[pos:])
        apply_run_style(run, bold=base_bold, color=base_color, size=base_size)


def style_paragraph(paragraph, *, after=6, before=0, line=1.10) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = line


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.08

    if "Caption" in styles:
        caption = styles["Caption"]
    else:
        caption = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(8)

    if "TOC Entry" not in styles:
        toc = styles.add_style("TOC Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc = styles["TOC Entry"]
    toc.font.name = "Calibri"
    toc.font.size = Pt(10)
    toc.font.color.rgb = RGBColor.from_string(INK)
    toc.paragraph_format.left_indent = Inches(0.15)
    toc.paragraph_format.space_after = Pt(3)


def configure_section(section, landscape=False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = LETTER_H
        section.page_height = LETTER_W
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = LETTER_W
        section.page_height = LETTER_H
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)


def setup_header_footer(section, first=False) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not first:
        add_inline(header_p, "EZEDU AI · BÁO CÁO PHÂN TÍCH NGHIỆP VỤ", base_bold=True, base_color=NAVY, base_size=8)
        p_pr = header_p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), LINE)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not first:
        add_inline(footer_p, "Báo cáo khảo sát & phân tích · Trang ", base_color=MUTED, base_size=8)
        add_page_field(footer_p)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, landscape=False)
    section.different_first_page_header_footer = True
    setup_header_footer(section, first=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(70)
    p.add_run("\n")

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag.add_run("SYSTEM ANALYSIS · UML · CASE STUDIO 2")
    apply_run_style(tag_run, bold=True, color=BLUE, size=10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run("BÁO CÁO KHẢO SÁT\nVÀ PHÂN TÍCH NGHIỆP VỤ")
    apply_run_style(title_run, bold=True, color=NAVY, size=26)

    product = doc.add_paragraph()
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product.paragraph_format.space_after = Pt(20)
    product_run = product.add_run("HỆ THỐNG EZEDU AI")
    apply_run_style(product_run, bold=True, color=BLUE, size=22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.left_indent = Inches(0.6)
    subtitle.paragraph_format.right_indent = Inches(0.6)
    subtitle.paragraph_format.space_after = Pt(36)
    add_inline(
        subtitle,
        "Hệ thống sinh câu hỏi đánh giá năng lực tự động từ học liệu điện tử bằng mô hình ngôn ngữ lớn",
        base_color=MUTED,
        base_size=13,
    )

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = band.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=180, bottom=180, start=260, end=260)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(cp, "BỘ BÀN GIAO PHÂN TÍCH HIỆN TRẠNG VÀ THIẾT KẾ LOGIC", base_bold=True, base_color=WHITE, base_size=11)

    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    values = [
        ("Phiên bản", "1.0"),
        ("Ngày khảo sát", "29/07/2026"),
        ("Phạm vi", "Frontend, backend, dữ liệu, AI, worker và tài liệu QA"),
        ("Chuẩn mô hình", "UML 2.x và ERD logic"),
        ("Công cụ đích", "CASE Studio 2 qua reverse engineering DDL MySQL"),
    ]
    for i, (label, value) in enumerate(values):
        set_cell_width(meta.cell(i, 0), 1800)
        set_cell_width(meta.cell(i, 1), 5400)
        set_cell_margins(meta.cell(i, 0))
        set_cell_margins(meta.cell(i, 1))
        set_cell_shading(meta.cell(i, 0), TABLE_HEADER)
        add_inline(meta.cell(i, 0).paragraphs[0], label, base_bold=True, base_color=NAVY, base_size=9)
        add_inline(meta.cell(i, 1).paragraphs[0], value, base_color=INK, base_size=9)

    bottom = doc.add_paragraph()
    bottom.paragraph_format.space_before = Pt(55)
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(bottom, "Tài liệu phục vụ phân tích, nghiệm thu và tiếp tục thiết kế hệ thống", base_color=MUTED, base_size=9)
    bottom.add_run().add_break(WD_BREAK.PAGE)


def add_static_toc(doc: Document, headings: list[tuple[int, str]]) -> None:
    heading = doc.add_paragraph("MỤC LỤC", style="Heading 1")
    add_bookmark(heading, "muc_luc", 1)
    doc.add_paragraph(
        "Các mục dưới đây liên kết đến nội dung tương ứng trong tài liệu.",
        style="Normal",
    )
    bookmark_id = 100
    for level, label in headings:
        if level > 3:
            continue
        p = doc.add_paragraph(style="TOC Entry")
        p.paragraph_format.left_indent = Inches(0.15 + (level - 2) * 0.25)
        anchor = f"sec_{bookmark_id}"
        add_hyperlink(p, label, anchor)
        bookmark_id += 1
    doc.add_page_break()


def table_col_widths(rows: list[list[str]], total: int) -> list[int]:
    cols = len(rows[0])
    scores = []
    for i in range(cols):
        content = [row[i] if i < len(row) else "" for row in rows]
        max_len = min(max((len(re.sub(r"[`*]", "", value)) for value in content), default=8), 60)
        scores.append(max(8, max_len))
    score_sum = sum(scores)
    widths = [max(800, int(total * score / score_sum)) for score in scores]
    scale = total / sum(widths)
    return [int(w * scale) for w in widths]


def add_markdown_table(doc: Document, rows: list[list[str]], landscape: bool) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = 13680 if landscape else 9360
    set_table_width(table, total)
    set_table_indent(table, 0 if landscape else 120)
    widths = table_col_widths(normalized, total)
    font_size = 7.3 if cols >= 6 else 8.0 if cols >= 5 else 8.5 if cols >= 4 else 9

    for r_idx, row in enumerate(normalized):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, after=0, line=1.0)
            add_inline(
                p,
                value,
                base_bold=(r_idx == 0),
                base_color=WHITE if r_idx == 0 else INK,
                base_size=font_size,
            )
            if r_idx == 0:
                set_cell_shading(cell, NAVY)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, PALE)
        if r_idx == 0:
            set_repeat_table_header(table.rows[0])
        table.rows[r_idx]._tr.get_or_add_trPr()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def switch_section(doc: Document, landscape: bool):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    separator = doc.paragraphs[-1]
    separator.paragraph_format.space_before = Pt(0)
    separator.paragraph_format.space_after = Pt(0)
    separator.paragraph_format.line_spacing = Pt(1)
    configure_section(section, landscape=landscape)
    setup_header_footer(section, first=False)
    return section


def add_diagram_content(doc: Document, image_path: Path, alt: str, caption_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    with Image.open(image_path) as image:
        aspect = image.width / image.height
    width_inches = min(9.2, 6.0 * aspect)
    height_inches = width_inches / aspect
    inline = p.add_run().add_picture(
        str(image_path),
        width=Inches(width_inches),
        height=Inches(height_inches),
    )
    set_picture_alt_text(
        inline,
        alt,
        f"{alt} của EzEdu AI; sơ đồ đầy đủ có bản PNG, SVG và PlantUML trong thư mục diagrams.",
    )
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(cap, caption_text, base_color=MUTED, base_size=9)


def parse_markdown(doc: Document, lines: list[str]) -> None:
    paragraph_buffer: list[str] = []
    bookmark_id = 100
    in_landscape_table = False
    diagram_section_active = False
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        value = " ".join(part.strip() for part in paragraph_buffer).strip()
        if value:
            p = doc.add_paragraph()
            style_paragraph(p)
            add_inline(p, value)
        paragraph_buffer = []

    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped == "---":
            flush_paragraph()
            i += 1
            continue

        diagram_heading = re.match(r"^###\s+(5\.[1-5]\s+.+)$", stripped)
        if diagram_heading:
            flush_paragraph()
            j = i + 1
            description_parts: list[str] = []
            image_match = None
            while j < len(lines) and j <= i + 8:
                candidate = lines[j].strip()
                if not candidate:
                    j += 1
                    continue
                image_match = re.match(r"!\[(.+?)\]\((.+?)\)", candidate)
                if image_match:
                    break
                if candidate.startswith("#"):
                    image_match = None
                    break
                description_parts.append(candidate)
                j += 1
            if image_match:
                page_break_before = False
                if not diagram_section_active:
                    switch_section(doc, landscape=True)
                    diagram_section_active = True
                else:
                    page_break_before = True
                label = diagram_heading.group(1)
                heading_p = doc.add_paragraph(style="Heading 2")
                heading_p.paragraph_format.page_break_before = page_break_before
                add_inline(heading_p, label, base_bold=True)
                add_bookmark(heading_p, f"sec_{bookmark_id}", bookmark_id)
                bookmark_id += 1
                if description_parts:
                    description = doc.add_paragraph()
                    style_paragraph(description, after=5)
                    add_inline(description, " ".join(description_parts))
                alt = image_match.group(1)
                image_path = (SOURCE.parent / image_match.group(2)).resolve()
                caption = {
                    "Use Case Diagram": "Hình 1. Use Case Diagram tổng thể của EzEdu AI.",
                    "Activity Diagram": "Hình 2. Activity Diagram vòng đời tài liệu và các hướng khai thác.",
                    "Sequence Diagram": "Hình 3. Sequence Diagram luồng hỏi đáp RAG theo tài liệu.",
                    "Class Diagram": "Hình 4. Class Diagram mô hình miền nghiệp vụ.",
                    "ERD": "Hình 5. ERD logic dùng cho CASE Studio 2; DDL đi kèm chứa mô hình đầy đủ.",
                }.get(alt, alt)
                add_diagram_content(doc, image_path, alt, caption)
                i = j + 1
                continue

        risk_group = re.match(r"^##\s+7\.\s+RỦI RO, KHOẢNG TRỐNG VÀ KIẾN NGHỊ$", stripped)
        if risk_group:
            flush_paragraph()
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            subheading = lines[j].strip() if j < len(lines) else ""
            if re.match(r"^###\s+7\.1\s+Sổ rủi ro$", subheading, re.IGNORECASE):
                k = j + 1
                while k < len(lines) and not lines[k].strip():
                    k += 1
                table_lines = []
                table_end = k
                while table_end < len(lines):
                    candidate = lines[table_end].strip()
                    if not (candidate.startswith("|") and candidate.endswith("|")):
                        break
                    table_lines.append([cell.strip() for cell in candidate[1:-1].split("|")])
                    table_end += 1
                if table_lines:
                    if len(table_lines) > 1 and all(
                        re.fullmatch(r":?-{3,}:?", cell) for cell in table_lines[1]
                    ):
                        table_lines.pop(1)
                    page_break_before = False
                    if diagram_section_active:
                        page_break_before = True
                    else:
                        switch_section(doc, landscape=True)
                        diagram_section_active = True
                    heading_p = doc.add_paragraph(style="Heading 1")
                    heading_p.paragraph_format.page_break_before = page_break_before
                    add_inline(heading_p, re.sub(r"^##\s+", "", stripped), base_bold=True)
                    add_bookmark(heading_p, f"sec_{bookmark_id}", bookmark_id)
                    bookmark_id += 1
                    sub_p = doc.add_paragraph(style="Heading 2")
                    add_inline(sub_p, re.sub(r"^###\s+", "", subheading), base_bold=True)
                    add_bookmark(sub_p, f"sec_{bookmark_id}", bookmark_id)
                    bookmark_id += 1
                    add_markdown_table(doc, table_lines, landscape=True)
                    i = table_end
                    continue

        heading_with_wide_table = re.match(
            r"^##\s+6\.\s+MA TRẬN TRUY VẾT$",
            stripped,
            re.IGNORECASE,
        )
        if heading_with_wide_table:
            flush_paragraph()
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            table_lines = []
            k = j
            while k < len(lines):
                candidate = lines[k].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append([cell.strip() for cell in candidate[1:-1].split("|")])
                k += 1
            if table_lines:
                if len(table_lines) > 1 and all(
                    re.fullmatch(r":?-{3,}:?", cell) for cell in table_lines[1]
                ):
                    table_lines.pop(1)
                page_break_before = False
                if diagram_section_active:
                    page_break_before = True
                else:
                    switch_section(doc, landscape=True)
                level = 2 if stripped.startswith("## ") else 3
                label = re.sub(r"^#{2,3}\s+", "", stripped)
                heading_p = doc.add_paragraph(style="Heading 1" if level == 2 else "Heading 2")
                heading_p.paragraph_format.page_break_before = page_break_before
                add_inline(heading_p, label, base_bold=True)
                add_bookmark(heading_p, f"sec_{bookmark_id}", bookmark_id)
                bookmark_id += 1
                add_markdown_table(doc, table_lines, landscape=True)
                diagram_section_active = True
                i = k
                continue

        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_paragraph()
            alt = image_match.group(1)
            image_path = (SOURCE.parent / image_match.group(2)).resolve()
            caption = {
                "Use Case Diagram": "Hình 1. Use Case Diagram tổng thể của EzEdu AI.",
                "Activity Diagram": "Hình 2. Activity Diagram vòng đời tài liệu và các hướng khai thác.",
                "Sequence Diagram": "Hình 3. Sequence Diagram luồng hỏi đáp RAG theo tài liệu.",
                "Class Diagram": "Hình 4. Class Diagram mô hình miền nghiệp vụ.",
                "ERD": "Hình 5. ERD logic dùng cho CASE Studio 2; DDL đi kèm chứa mô hình đầy đủ.",
            }.get(alt, alt)
            switch_section(doc, landscape=True)
            add_diagram_content(doc, image_path, alt, caption)
            switch_section(doc, landscape=False)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            label = heading_match.group(2).strip()
            if level == 1:
                i += 1
                continue
            style_name = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            p = doc.add_paragraph(style=style_name)
            add_inline(p, label, base_bold=True)
            if level <= 3:
                add_bookmark(p, f"sec_{bookmark_id}", bookmark_id)
                bookmark_id += 1
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_lines = []
            while i < len(lines):
                candidate = lines[i].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                cells = [cell.strip() for cell in candidate[1:-1].split("|")]
                table_lines.append(cells)
                i += 1
            if len(table_lines) > 1 and all(re.fullmatch(r":?-{3,}:?", c) for c in table_lines[1]):
                table_lines.pop(1)
            col_count = len(table_lines[0]) if table_lines else 0
            if col_count >= 6:
                switch_section(doc, landscape=True)
                in_landscape_table = True
            add_markdown_table(doc, table_lines, landscape=in_landscape_table)
            if in_landscape_table:
                switch_section(doc, landscape=False)
                in_landscape_table = False
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet_match.group(1))
            i += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, number_match.group(1))
            i += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(6)
            add_inline(p, stripped[2:], base_color=MUTED)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()


def add_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Báo cáo khảo sát và phân tích nghiệp vụ hệ thống EzEdu AI"
    props.subject = "Khảo sát hiện trạng, yêu cầu chức năng/phi chức năng và UML/ERD"
    props.author = "Codex – phân tích từ mã nguồn EzEdu AI"
    props.keywords = "EzEdu AI, business analysis, UML, ERD, CASE Studio 2"
    props.comments = "Bản 1.0 – 29/07/2026"


def main() -> int:
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    first_rule = next((idx for idx, line in enumerate(source_lines) if line.strip() == "---"), 7)
    content_lines = source_lines[first_rule + 1 :]
    headings = []
    for line in content_lines:
        match = re.match(r"^(#{2,3})\s+(.+)$", line.strip())
        if match:
            headings.append((len(match.group(1)), match.group(2).strip()))

    doc = Document()
    setup_styles(doc)
    add_document_properties(doc)
    add_cover(doc)
    add_static_toc(doc, headings)
    parse_markdown(doc, content_lines)

    for section in doc.sections:
        if not section.header.is_linked_to_previous and not section.footer.is_linked_to_previous:
            continue
        setup_header_footer(section, first=False)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Sections: {len(doc.sections)}; paragraphs: {len(doc.paragraphs)}; tables: {len(doc.tables)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
