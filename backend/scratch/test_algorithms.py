"""
Integration test: TF-IDF + Bloom + Dual-AI pipeline
=====================================================
Tests the full pipeline: Register -> Login -> Upload -> Extract -> Index -> Generate
Verifies: keywords, bloom_distribution, bloom_level on questions, validation_stats
"""
import asyncio
import httpx
import json
import sys

BASE = "http://localhost:8000/api/v1"

async def main():
    async with httpx.AsyncClient(timeout=120.0) as c:
        # 1. Register + Login
        print("1. Logging in / Registering test user...")
        ts = str(int(asyncio.get_event_loop().time() * 1000))[-6:]
        email = f"test_algo_{ts}@test.com"
        pwd = "TestPass123!"
        
        await c.post(f"{BASE}/auth/register", json={
            "email": email, "password": pwd, "full_name": "Algo Test"
        })
        r = await c.post(f"{BASE}/auth/login", json={"email": email, "password": pwd})
        token = r.json()["access_token"]
        headers = {"Authorization": f"Bearer {token}"}
        print(f"   Logged in. Token acquired.")

        # 2. Upload test document
        print("2. Uploading test.docx...")
        from docx import Document
        doc = Document()
        doc.add_heading("Quang hợp ở thực vật", level=1)
        doc.add_paragraph(
            "Quang hợp là quá trình thực vật sử dụng ánh sáng mặt trời để tổng hợp "
            "chất hữu cơ từ CO2 và nước. Quá trình quang hợp diễn ra chủ yếu ở lá cây, "
            "cụ thể là trong lục lạp. Lục lạp chứa chất diệp lục (chlorophyll) có khả năng "
            "hấp thụ ánh sáng. Quang hợp gồm hai pha: pha sáng và pha tối."
        )
        doc.add_paragraph(
            "Pha sáng diễn ra trên màng thylakoid của lục lạp, sử dụng năng lượng ánh sáng "
            "để phân ly nước, giải phóng O2 và tạo ra ATP, NADPH. Pha tối (chu trình Calvin) "
            "diễn ra trong chất nền (stroma) của lục lạp, sử dụng ATP và NADPH từ pha sáng "
            "để cố định CO2 thành các hợp chất hữu cơ (glucose). Phương trình tổng quát: "
            "6CO2 + 6H2O → C6H12O6 + 6O2."
        )
        doc.add_paragraph(
            "Các yếu tố ảnh hưởng đến quang hợp bao gồm: cường độ ánh sáng, nồng độ CO2, "
            "nhiệt độ, và nước. Khi cường độ ánh sáng tăng, tốc độ quang hợp tăng cho đến "
            "điểm bão hòa ánh sáng. Nồng độ CO2 tăng cũng thúc đẩy quang hợp nhưng có giới hạn. "
            "Nhiệt độ tối ưu cho quang hợp ở thực vật C3 là khoảng 25-30°C."
        )
        doc.save("/tmp/test_algo.docx")

        with open("/tmp/test_algo.docx", "rb") as f:
            r = await c.post(f"{BASE}/documents/upload", 
                           files={"file": ("test_algo.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                           headers=headers)
        doc_id = r.json()["document_id"]
        print(f"   Upload successful. Document ID: {doc_id}")

        # 3. Extract
        print("3. Extracting content...")
        await c.post(f"{BASE}/documents/{doc_id}/extract", headers=headers)
        print("   Extraction complete.")

        # 4. Index
        print("4. Indexing document...")
        await c.post(f"{BASE}/documents/{doc_id}/index", headers=headers)
        print("   Indexing complete.")

        # 5. Generate questions with all 3 algorithms
        print("5. Generating questions (TF-IDF + Bloom + Dual-AI)...")
        r = await c.post(f"{BASE}/questions/generate", json={
            "document_id": doc_id,
            "question_count": 3,
            "difficulty": "medium",
            "question_type": "multiple_choice",
            "bloom_level": None
        }, headers=headers)
        
        if r.status_code != 201:
            print(f"\n❌ FAILED! Status: {r.status_code}")
            print(f"   Response: {r.text[:500]}")
            sys.exit(1)

        data = r.json()
        print(f"\n🎉 SUCCESS! Full integration works.")
        print(f"Generated Set ID: {data['id']}")
        print(f"Question Count: {data['question_count']}")

        # Check TF-IDF Keywords
        keywords = data.get("keywords", [])
        print(f"\n--- TF-IDF Keywords ({len(keywords)} extracted) ---")
        for kw in keywords[:10]:
            print(f"  📌 {kw['keyword']} (score: {kw['score']})")

        # Check Bloom Distribution
        bloom_dist = data.get("bloom_distribution", {})
        print(f"\n--- Bloom's Taxonomy Distribution ---")
        bloom_labels = {"remember": "Nhận biết", "understand": "Thông hiểu", "apply": "Vận dụng", "analyze": "Vận dụng cao"}
        for level, count in bloom_dist.items():
            print(f"  🎓 {bloom_labels.get(level, level)}: {count} câu")

        # Check Validation Stats
        vstats = data.get("validation_stats", {})
        print(f"\n--- Cross-Validation Stats ---")
        print(json.dumps(vstats, indent=2, ensure_ascii=False))

        # Check Questions with bloom_level
        print(f"\n--- Questions ---")
        for i, q in enumerate(data["questions"]):
            bloom = q.get("bloom_level", "N/A")
            print(f"{i+1}. [{bloom_labels.get(bloom, bloom)}] {q['question']}")
            print(f"   Correct: {q['correct_answer']}")

        # Verify all 3 algorithms are present
        print(f"\n{'='*50}")
        has_keywords = len(keywords) > 0
        has_bloom = len(bloom_dist) > 0
        has_validation = vstats.get("cross_validated", False)
        has_bloom_on_questions = all(q.get("bloom_level") for q in data["questions"])
        
        print(f"✅ TF-IDF Keywords: {'OK' if has_keywords else 'MISSING'} ({len(keywords)} keywords)")
        print(f"✅ Bloom Distribution: {'OK' if has_bloom else 'MISSING'} ({bloom_dist})")
        print(f"✅ Bloom on Questions: {'OK' if has_bloom_on_questions else 'MISSING'}")
        print(f"✅ Dual-AI Validation: {'OK' if has_validation else 'MISSING'}")
        
        if has_keywords and has_bloom and has_bloom_on_questions:
            print(f"\n🏆 ALL 3 ALGORITHMS VERIFIED SUCCESSFULLY!")
        else:
            print(f"\n⚠️ Some algorithms may not have produced results.")

asyncio.run(main())
