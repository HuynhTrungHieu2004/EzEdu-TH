import io
import os
import re
from datetime import datetime
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.core.config import settings

EXPORT_TITLE = "BỘ CÂU HỎI ĐÁNH GIÁ NĂNG LỰC"


def _format_datetime(value) -> str:
    if isinstance(value, datetime):
        return value.strftime("%H:%M %d/%m/%Y")
    if value:
        return str(value)
    return "Không xác định"


def _difficulty_label(value: str | None) -> str:
    mapping = {
        "easy": "Dễ",
        "medium": "Trung bình",
        "hard": "Khó",
    }
    return mapping.get((value or "").lower(), value or "Không xác định")


def _question_type_label(value: str | None) -> str:
    mapping = {
        "multiple_choice": "Trắc nghiệm",
        "true_false": "Đúng/Sai",
        "short_answer": "Tự luận ngắn",
    }
    return mapping.get((value or "").lower(), value or "Không xác định")


def _system_name() -> str:
    return settings.PROJECT_NAME or "Hệ thống sinh câu hỏi tự động"


def _question_options(question: dict) -> list[tuple[str, str]]:
    options = question.get("options")
    if isinstance(options, dict):
        return [(str(key), str(value)) for key, value in options.items()]
    return []


def _register_pdf_font() -> str:
    font_candidates = [
        ("ArialUnicode", "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        ("Arial", "/System/Library/Fonts/Supplemental/Arial.ttf"),
        ("Arial", "/Library/Fonts/Arial.ttf"),
        ("Helvetica", None),
        ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ("NotoSans", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
    ]

    for font_name, font_path in font_candidates:
        if font_path is None:
            return font_name
        if os.path.exists(font_path):
            try:
                if font_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                return font_name
            except Exception:
                continue

    return "Helvetica"


def build_export_filename(question_set: dict, extension: str) -> str:
    raw_name = question_set.get("document_name") or "bo_cau_hoi"
    slug = re.sub(r"[^A-Za-z0-9_-]+", "_", raw_name).strip("_") or "bo_cau_hoi"
    return f"{slug}_question_set.{extension}"


def export_question_set_to_docx(question_set: dict) -> io.BytesIO:
    document = Document()
    document.core_properties.title = EXPORT_TITLE
    document.core_properties.subject = question_set.get("document_name", "Tài liệu không tên")
    document.core_properties.author = _system_name()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    title = document.add_heading(EXPORT_TITLE, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(_system_name()).italic = True

    document.add_paragraph(f"Tài liệu: {question_set.get('document_name', 'Tài liệu không tên')}")
    document.add_paragraph(f"Ngày tạo bộ câu hỏi: {_format_datetime(question_set.get('created_at'))}")
    document.add_paragraph(f"Loại câu hỏi: {_question_type_label(question_set.get('question_type'))}")
    document.add_paragraph(f"Độ khó: {_difficulty_label(question_set.get('difficulty'))}")
    document.add_paragraph("")

    for index, question in enumerate(question_set.get("questions", []), start=1):
        document.add_paragraph(f"Câu {index}: {question.get('question', '')}", style="List Number")

        for option_key, option_value in _question_options(question):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.add_run(f"{option_key}. {option_value}")

        answer_paragraph = document.add_paragraph()
        answer_paragraph.paragraph_format.left_indent = Pt(18)
        answer_paragraph.add_run("Đáp án đúng: ").bold = True
        answer_paragraph.add_run(str(question.get("correct_answer", "Không xác định")))

        explanation_paragraph = document.add_paragraph()
        explanation_paragraph.paragraph_format.left_indent = Pt(18)
        explanation_paragraph.add_run("Giải thích: ").bold = True
        explanation_paragraph.add_run(str(question.get("explanation", "Không có giải thích.")))

        document.add_paragraph("")

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


def export_question_set_to_pdf(question_set: dict) -> io.BytesIO:
    font_name = _register_pdf_font()
    file_stream = io.BytesIO()
    pdf = SimpleDocTemplate(
        file_stream,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
        title=EXPORT_TITLE,
        author=_system_name(),
    )

    base_styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=base_styles["Heading1"],
        fontName=font_name,
        fontSize=17,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "ExportSubtitle",
        parent=base_styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "ExportMeta",
        parent=base_styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    question_style = ParagraphStyle(
        "ExportQuestion",
        parent=base_styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=16,
        spaceBefore=8,
        spaceAfter=6,
    )
    option_style = ParagraphStyle(
        "ExportOption",
        parent=base_styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        leftIndent=18,
        spaceAfter=2,
    )
    answer_style = ParagraphStyle(
        "ExportAnswer",
        parent=base_styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        leftIndent=18,
        spaceAfter=3,
    )

    story = [
        Paragraph(f"<b>{escape(EXPORT_TITLE)}</b>", title_style),
        Paragraph(escape(_system_name()), subtitle_style),
        Paragraph(f"<b>Tài liệu:</b> {escape(question_set.get('document_name', 'Tài liệu không tên'))}", meta_style),
        Paragraph(f"<b>Ngày tạo bộ câu hỏi:</b> {escape(_format_datetime(question_set.get('created_at')))}", meta_style),
        Paragraph(f"<b>Loại câu hỏi:</b> {escape(_question_type_label(question_set.get('question_type')))}", meta_style),
        Paragraph(f"<b>Độ khó:</b> {escape(_difficulty_label(question_set.get('difficulty')))}", meta_style),
        Spacer(1, 10),
    ]

    for index, question in enumerate(question_set.get("questions", []), start=1):
        story.append(
            Paragraph(
                f"<b>Câu {index}:</b> {escape(str(question.get('question', '')))}",
                question_style,
            )
        )

        for option_key, option_value in _question_options(question):
            story.append(
                Paragraph(
                    f"{escape(option_key)}. {escape(option_value)}",
                    option_style,
                )
            )

        story.append(
            Paragraph(
                f"<b>Đáp án đúng:</b> {escape(str(question.get('correct_answer', 'Không xác định')))}",
                answer_style,
            )
        )
        story.append(
            Paragraph(
                f"<b>Giải thích:</b> {escape(str(question.get('explanation', 'Không có giải thích.')))}",
                option_style,
            )
        )
        story.append(Spacer(1, 8))

    pdf.build(story)
    file_stream.seek(0)
    return file_stream
